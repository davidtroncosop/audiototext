import streamlit as st
import os
from groq import Groq
from docx import Document
from io import BytesIO
from dotenv import load_dotenv

# Carga las variables de entorno desde el archivo .env
load_dotenv()

# Configura la clave API
api_key = os.getenv("GROQ_API_KEY")

# Inicializa el cliente de Groq
client = Groq(api_key=api_key)

st.title("Transcripción de audio")
st.subheader("by David Troncoso")

# Subir archivo de audio
uploaded_file = st.file_uploader("Sube un archivo de audio (m4a, ogg, mp3)", type=["m4a", "ogg", "mp3"])

if uploaded_file is not None:
    st.write("Archivo cargado exitosamente.")

    # Transcribir el archivo de audio
    transcription = client.audio.transcriptions.create(
        file=(uploaded_file.name, uploaded_file.read()),  # Archivo de audio requerido
        model="distil-whisper-large-v3-en",  # Modelo para la transcripción
        prompt="Specify context or spelling",  # Opcional
        response_format="json",  # Opcional
        temperature=0.0  # Opcional
    )
    
    # Muestra el texto transcrito
    transcribed_text = transcription.text


    # Define el prompt para estructurar la información transcrita
    prompt = f"""
    Tienes el siguiente texto transcrito:

    {transcribed_text}

Por favor traduce el texto a español.     
"""

    # Generar la información estructurada con el prompt
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="llama3-8b-8192",
    )
    
    structured_text = chat_completion.choices[0].message.content

    # Mostrar la información estructurada generada
    st.subheader("Información Estructurada")
    st.write(structured_text)

    # Crear un archivo .docx con la transcripción y la información estructurada
    doc = Document()
    doc.add_heading('Transcripción y Resumen de Audio', 0)

    doc.add_paragraph('by David Troncoso')

    doc.add_heading('Texto Transcrito', level=1)
    doc.add_paragraph(transcribed_text)

    doc.add_heading('Información Estructurada', level=1)
    doc.add_paragraph(structured_text)

    # Guardar el documento en un buffer de memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Ofrecer el archivo .docx para descarga
    st.download_button(
        label="Descargar el documento",
        data=buffer,
        file_name="transcripcion_y_resumen.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
